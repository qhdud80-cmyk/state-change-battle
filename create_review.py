from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
import os

doc = Document()

# Page setup - A4
section = doc.sections[0]
section.page_height = Cm(29.7)
section.page_width = Cm(21.0)
section.top_margin = Cm(1.5)
section.bottom_margin = Cm(1.5)
section.left_margin = Cm(2.0)
section.right_margin = Cm(2.0)

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = '맑은 고딕'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

# ---- Title ----
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('애니메이션 「룩백 (Look Back)」 감상문')
run.bold = True
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('─  중학교 1학년  ─')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()

# ---- Student Info Table ----
table = doc.add_table(rows=2, cols=4)
table.style = 'Table Grid'
table.alignment = WD_ALIGN_PARAGRAPH.CENTER

cells_text = [('학년 반', '1학년 (  )반', '번호', '(  )번'),
              ('이름', '', '작성일', '2025년   월   일')]

for i, row_data in enumerate(cells_text):
    for j, text in enumerate(row_data):
        cell = table.cell(i, j)
        cell.text = text
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.size = Pt(11)
            run.font.name = '맑은 고딕'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

# Set first column width
for row in table.rows:
    row.cells[0].width = Cm(2.5)

doc.add_paragraph()

# ---- Guide ----
p = doc.add_paragraph()
run = p.add_run('■  다음 내용을 참고하여 감상문을 작성해 봅시다.')
run.bold = True
run.font.size = Pt(11)
run.font.name = '맑은 고딕'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

doc.add_paragraph()

# ---- Question 1 ----
def add_question(num, title, prompt, space_cm=2.0):
    p = doc.add_paragraph()
    run = p.add_run(f'{num}. {title}')
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = '맑은 고딕'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    
    p = doc.add_paragraph()
    run = p.add_run(prompt)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    run.font.name = '맑은 고딕'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    
    for _ in range(int(space_cm / 0.7)):
        doc.add_paragraph()

add_question(1, '인상 깊었던 장면',
    '• 영화에서 가장 기억에 남는 장면은 무엇인가요? 그 이유도 함께 적어 보세요.')
add_question(2, '후지노와 쿄모토의 우정',
    '• 그림이라는 공통된 관심사로 시작된 두 친구의 관계를 보면서 어떤 점이 가장 인상 깊었나요?')
add_question(3, '그림 그리기의 의미',
    '• 영화 속에서 그림 그리기는 두 주인공에게 어떤 의미였나요? 나에게도 비슷한 경험이 있나요?')
add_question(4, '「룩백」의 의미',
    '• 영화 제목 「룩백 (Look Back)」은 무슨 뜻일까요? 영화 내용과 연결지어 생각해 봅시다.')
add_question(5, '나의 다짐',
    '• 이 영화를 감상한 후, 나는 앞으로 어떤 마음가짐으로 나의 꿈이나 관심사를 키워 나가고 싶은가요?')

# ---- Signature ----
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p.add_run('위 학생은 위와 같이 감상문을 작성하였습니다.')
run.font.size = Pt(10)
run.font.name = '맑은 고딕'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

# Save as DOCX (한글에서 열 수 있음)
out_dir = os.path.join(os.environ['USERPROFILE'], 'Desktop')
output_path = os.path.join(out_dir, '룩백_감상문.docx')
doc.save(output_path)
print(f'저장 완료: {output_path}')
