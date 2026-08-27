from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

section = doc.sections[0]
section.page_height = Cm(29.7)
section.page_width = Cm(21.0)
section.top_margin = Cm(1.0)
section.bottom_margin = Cm(1.0)
section.left_margin = Cm(1.3)
section.right_margin = Cm(1.3)

style = doc.styles['Normal']
font = style.font
font.name = '맑은 고딕'
font.size = Pt(10)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
style.paragraph_format.space_after = Pt(1)
style.paragraph_format.space_before = Pt(1)

# Summer palette
SKY = '7DD3FC'       # 하늘색
OCEAN = '38BDF8'     # 바다색
SAND = 'FEF3C7'      # 모래색
CORAL = 'FCA5A5'     # 코랄
MINT = 'A7F3D0'      # 민트
WHITE = 'FFFFFF'
DARK = '0C4A6E'      # 진한 바다색
GRAY = '94A3B8'
SOFT_BLUE = 'E0F2FE'

def set_cell_font(cell, name='맑은 고딕', size=Pt(9), bold=False, align=WD_ALIGN_PARAGRAPH.CENTER, color=None):
    for p in cell.paragraphs:
        p.alignment = align
        for run in p.runs:
            run.font.name = name
            run.font.size = size
            run.bold = bold
            run.element.rPr.rFonts.set(qn('w:eastAsia'), name)
            if color:
                run.font.color.rgb = RGBColor(*color)

def shade(cell, color):
    s = OxmlElement('w:shd')
    s.set(qn('w:fill'), color)
    s.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(s)

def write_heading(text, size=Pt(16)):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.size = size
    run.font.color.rgb = RGBColor(0x0C, 0x4A, 0x6E)
    run.font.name = '맑은 고딕'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

def write_sub(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f'~ {text}')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
    run.font.name = '맑은 고딕'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

def make_table(rows, cols, data, col_widths=None):
    t = doc.add_table(rows=rows, cols=cols)
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row_data in enumerate(data):
        for j, cell_text in enumerate(row_data):
            t.cell(i, j).text = str(cell_text)
            sz = Pt(9) if i == 0 else Pt(8)
            set_cell_font(t.cell(i, j), size=sz, bold=(i == 0))
            if i == 0:
                shade(t.cell(i, j), SKY)
            elif i % 2 == 0:
                shade(t.cell(i, j), SOFT_BLUE)
    if col_widths:
        for row in t.rows:
            for j, w in enumerate(col_widths):
                if j < len(row.cells):
                    row.cells[j].width = Cm(w)
    return t

def blank(n=1):
    for _ in range(n):
        doc.add_paragraph().paragraph_format.space_after = Pt(0)

# ============================================================
# COVER
# ============================================================
for _ in range(3):
    blank()

# Decorative title with wave
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('~ ~ ~ ~ ~ ~ ~ ~ ~')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x56, 0xBC, 0xE8)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('여름 방학')
run.bold = True
run.font.size = Pt(34)
run.font.color.rgb = RGBColor(0x0C, 0x4A, 0x6E)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('S T U D Y   P L A N N E R')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x38, 0xBD, 0xF8)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('~ ~ ~ ~ ~ ~ ~ ~ ~')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x56, 0xBC, 0xE8)

blank(3)

info_t = doc.add_table(rows=5, cols=2)
info_t.alignment = WD_TABLE_ALIGNMENT.CENTER
info_t.style = 'Table Grid'
cover_items = [
    ('학년 / 반', '중학교 1학년 (  )반'),
    ('이  름', ''),
    ('방학 기간', '2025년  월  일  ~   월  일'),
    ('이번 방학 목표', ''),
    ('다짐 한마디', ''),
]
for i, (k, v) in enumerate(cover_items):
    info_t.cell(i, 0).text = k
    set_cell_font(info_t.cell(i, 0), size=Pt(11), bold=True, color=(0x0C, 0x4A, 0x6E))
    info_t.cell(i, 1).text = v
    set_cell_font(info_t.cell(i, 1), size=Pt(11))
    shade(info_t.cell(i, 0), SKY)
for row in info_t.rows:
    row.cells[0].width = Cm(3.5)

blank(2)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('시원한 바람과 함께, 알찬 방학을 만들어 봐요!')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

# ============================================================
# 과목별 목표
# ============================================================
doc.add_page_break()

write_heading('1. 과목별 목표 세우기')

subjects = ['국어', '영어', '수학', '사회', '과학', '기타']
goal_data = [['과목', '현재 실력', '방학 목표', '실천할 것', '✔']]
for s in subjects:
    goal_data.append([s, '', '', '', '□'])
make_table(len(goal_data), 5, goal_data, col_widths=[2.0, 2.5, 3.5, 6.5, 1.0])

blank(2)
write_heading('하루 공부 시간 목표')
blank()
time_data = [
    ['', '평일 (월~금)', '주말 (토/일)'],
    ['목표 시간', '_____ 시간', '_____ 시간'],
    ['집중 시간대', '_____ 시 ~ _____ 시', '_____ 시 ~ _____ 시'],
]
make_table(3, 3, time_data, col_widths=[3.5, 5.0, 5.0])

# ============================================================
# 주간 계획표 x6
# ============================================================
doc.add_page_break()

day_labels = ['', '월', '화', '수', '목', '금', '토', '일']
time_slots = [
    '아침 (~8:30)', '오전① (8:30~10:00)',
    '오전② (10:00~12:00)', '점심 (12:00~13:00)',
    '오후① (13:00~15:00)', '오후② (15:00~17:00)',
    '저녁 (17:00~20:00)', '자유시간 (20:00~)'
]

for week in range(1, 7):
    write_heading(f'{week}주차  (  /  ~  /  )')
    t = doc.add_table(rows=len(time_slots)+1, cols=9)
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    for j, d in enumerate(day_labels):
        t.cell(0, j).text = d
        set_cell_font(t.cell(0, j), size=Pt(8), bold=True)
        shade(t.cell(0, j), SKY)

    t.cell(0, 8).text = '✨'
    set_cell_font(t.cell(0, 8), size=Pt(8), bold=True)
    shade(t.cell(0, 8), SAND)

    for i, slot in enumerate(time_slots):
        t.cell(i+1, 0).text = slot
        set_cell_font(t.cell(i+1, 0), size=Pt(7), bold=True)
        shade(t.cell(i+1, 0), SOFT_BLUE)
        for j in range(1, 8):
            t.cell(i+1, j).text = ''
            set_cell_font(t.cell(i+1, j), size=Pt(7))
        t.cell(i+1, 8).text = ''

    for row in t.rows:
        row.cells[0].width = Cm(2.3)
        for j in range(1, 8):
            row.cells[j].width = Cm(1.8)
        row.cells[8].width = Cm(1.5)

    blank()
    p = doc.add_paragraph()
    run = p.add_run(f'  ~ 한 일을 짧게 적어요 (예: 수학 p.20~35)')
    run.font.size = Pt(7)
    run.font.color.rgb = RGBColor(0xCA, 0xD5, 0xE0)
    run.font.name = '맑은 고딕'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

    if week < 6:
        doc.add_page_break()

# ============================================================
# 오답 노트
# ============================================================
doc.add_page_break()

write_heading('2. 문제 풀이 & 오답 노트')
write_sub('틀린 문제를 다시 풀어야 실력이 쌓여요')
blank()

odata = [['날짜', '과목', '문제집', '페이지', '맞음', '틀림', '복습']]
for _ in range(14):
    odata.append(['', '', '', '', '', '', ''])
make_table(len(odata), 7, odata, col_widths=[2.0, 1.5, 3.0, 2.0, 1.5, 1.5, 1.5])

blank(2)
write_heading('취약 유형 & 보충 계획')
blank()
weak_data = [['과목', '취약한 부분', '보충 방법', '✔']]
for _ in range(6):
    weak_data.append(['', '', '', '□'])
make_table(len(weak_data), 4, weak_data, col_widths=[1.5, 5.0, 6.0, 1.0])

# ============================================================
# 암기 체크리스트
# ============================================================
doc.add_page_break()

write_heading('3. 암기 체크리스트')
write_sub('반복이 기억을 만듭니다. 조금씩 꾸준히!')
blank()

memo_data = [['날짜', '영단어', '한자/어휘', '과학', '사회', '기타', '✔']]
for _ in range(20):
    memo_data.append(['', '____개', '____개', '____개', '____개', '', '□'])
make_table(len(memo_data), 7, memo_data, col_widths=[2.0, 2.0, 2.0, 1.8, 1.8, 2.5, 1.0])

# ============================================================
# 2학기 예습
# ============================================================
doc.add_page_break()

write_heading('4. 2학기 예습 계획')
write_sub('미리 준비하면 2학기가 더 쉬워져요!')
blank()

prep_data = [['과목', '1학기 복습', '2학기 미리보기', '참고서/자료']]
for s in ['국어', '영어', '수학', '사회', '과학']:
    prep_data.append([s, '', '', ''])
make_table(len(prep_data), 4, prep_data, col_widths=[1.5, 4.5, 4.5, 4.0])

blank(2)
write_heading('2학기 목표 성적')
blank()
score_data = [
    ['', '국어', '영어', '수학', '사회', '과학'],
    ['목표', '', '', '', '', ''],
    ['각오', '', '', '', '', ''],
]
make_table(3, 6, score_data, col_widths=[1.5, 2.5, 2.5, 2.5, 2.5, 2.5])

# ============================================================
# 활동 체크리스트
# ============================================================
doc.add_page_break()

write_heading('5. 방학 도전! 활동 체크리스트')
write_sub('다양하게 도전하고 스스로 칭찬해요')
blank()

check_data = [['', '활동', '✔ (3번)']]
checks = [
    ['공부', '하루 1시간 이상 집중 공부', '□ □ □'],
    ['공부', '틀린 문제 복습하기', '□ □ □'],
    ['공부', '영어 단어 10개 이상 외우기', '□ □ □'],
    ['공부', '2학기 교재 미리 훑어보기', '□ □ □'],
    ['독서', '하루 20분 이상 책 읽기', '□ □ □'],
    ['독서', '읽은 책 짧은 감상 남기기', '□ □ □'],
    ['운동', '매일 스트레칭/산책', '□ □ □'],
    ['운동', '주 3회 이상 신나게 운동', '□ □ □'],
    ['취미', '새로운 취미 도전하기', '□ □ □'],
    ['취미', '꾸준히 하는 나만의 취미', '□ □ □'],
    ['생활', '아침 8시 전 기상', '□ □ □'],
    ['생활', '스마트폰 2시간 이하', '□ □ □'],
    ['생활', '책상 정리 / 방 청소', '□ □ □'],
    ['도전', '가족을 위해 요리하기', '□ □ □'],
    ['도전', '일기 7일 이상 쓰기', '□ □ □'],
    ['도전', '영화/다큐 보고 느낌 정리', '□ □ □'],
]
make_table(len(checks)+1, 3, check_data + checks, col_widths=[2.0, 9.0, 3.5])

# ============================================================
# 집중 타임 기록
# ============================================================
doc.add_page_break()

write_heading('6. 나의 집중 타임 기록')
write_sub('뽀모도로 타이머와 함께 집중력을 키워요')
blank()

focus_data = [['날짜', '과목', '시작', '종료', '공부시간', '집중★', '비고']]
for _ in range(22):
    focus_data.append(['', '', '__:__', '__:__', '__분', '', ''])
make_table(len(focus_data), 7, focus_data, col_widths=[1.8, 1.5, 1.5, 1.5, 1.5, 1.5, 3.0])

blank(2)
write_heading('주간 총 공부 시간')
blank()
daily_data = [['', '월', '화', '수', '목', '금', '토', '일', '합계']]
daily_data.append(['시간', '', '', '', '', '', '', '', ''])
make_table(2, 9, daily_data, col_widths=[1.5, 1.7, 1.7, 1.7, 1.7, 1.7, 1.7, 1.7, 1.7])

# ============================================================
# 매주 돌아보기
# ============================================================
doc.add_page_break()

write_heading('7. 매주 돌아보기')

for week in range(1, 7):
    write_sub(f'{week}주차를 돌아보며')
    rv_t = doc.add_table(rows=5, cols=2)
    rv_t.style = 'Table Grid'
    rv_t.alignment = WD_TABLE_ALIGNMENT.CENTER
    items = [
        ('이번 주 공부 시간', '약 ______시간'),
        ('가장 잘한 점', ''),
        ('아쉬운 점 / 개선할 점', ''),
        ('다음 주 목표', ''),
        ('스스로 ★ 점수', '★  ★  ★  ★  ★  (   점)'),
    ]
    for i, (k, v) in enumerate(items):
        rv_t.cell(i, 0).text = k
        set_cell_font(rv_t.cell(i, 0), size=Pt(9), bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, color=(0x0C, 0x4A, 0x6E))
        shade(rv_t.cell(i, 0), SKY)
        rv_t.cell(i, 1).text = v
        set_cell_font(rv_t.cell(i, 1), size=Pt(9), align=WD_ALIGN_PARAGRAPH.LEFT)
    for row in rv_t.rows:
        row.cells[0].width = Cm(4.5)
        row.cells[1].width = Cm(12.0)
    blank(2)

# ============================================================
# 마무리
# ============================================================
doc.add_page_break()

for _ in range(3):
    blank()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('방학을 마치며')
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x0C, 0x4A, 0x6E)

blank()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('스스로를 돌아보고, 한 걸음 더 성장한 나를 만나요')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

blank(2)

final = [
    ['이번 방학, 나는 총 ____시간 공부했어요!'],
    ['가장 잘한 일은?'],
    ['아쉬웠던 점은?'],
    ['2학기 꼭 실천할 것은?'],
    ['나에게 보내는 응원 메시지'],
]
ft = doc.add_table(rows=5, cols=1)
ft.style = 'Table Grid'
ft.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (txt,) in enumerate(final):
    ft.cell(i, 0).text = txt
    set_cell_font(ft.cell(i, 0), size=Pt(11), bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
    shade(ft.cell(i, 0), SOFT_BLUE) if i % 2 == 0 else None
for row in ft.rows:
    row.cells[0].width = Cm(17.0)

blank(3)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('파도처럼 밀려오는 방학,\n그 속에서 꾸준히 나아간 나는 최고야!')
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0x38, 0xBD, 0xF8)

blank()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('☀️  🏄  🌊  🧊  🍨')
run.font.size = Pt(18)

# Save
out_dir = os.path.join(os.environ['USERPROFILE'], 'Desktop')
output_path = os.path.join(out_dir, '여름방학_스터디플래너.docx')
doc.save(output_path)
print(f'저장 완료: {output_path}')
